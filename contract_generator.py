"""
Contract Generator - Fill Word template with KOC data from Lark Base
Version 2.0.0

PLACEHOLDER-BASED: Templates use {PLACEHOLDER} markers that get replaced with data.
Works with ANY template structure - no hardcoded paragraph/run indices.

Supported placeholders:
  {HO_TEN}              - Họ và Tên Bên B
  {DIA_CHI}             - Địa chỉ Bên B
  {MST}                 - Mã số thuế Bên B
  {SDT}                 - Số điện thoại Bên B
  {GMAIL}               - Gmail Bên B
  {CCCD}                - Số CCCD Bên B
  {CCCD_NGAY_CAP}       - Ngày cấp CCCD
  {CCCD_NOI_CAP}        - Nơi cấp CCCD
  {STK}                 - Số tài khoản Bên B
  {ID_KOC}              - ID KOC (tên kênh)
  {THUONG_HIEU}         - Thương hiệu
  {NGAY_DU_KIEN_AIR}    - Ngày dự kiến air
  {HOA_HONG_TU_NHIEN}   - % Hoa hồng tự nhiên (đã x100)
  {HOA_HONG_CHAY_ADS}   - % Hoa hồng chạy ads (đã x100)
  {SO_LUONG_CLIP}       - Số lượng clip
  {CHI_PHI}             - Chi phí 1 clip (formatted: 300.000)
  {THANH_TIEN}          - Thành tiền (formatted: 300.000)
  {THUE_TNCN}           - Thuế TNCN (formatted: 30.000)
  {TONG_SAU_THUE}       - Tổng giá trị sau thuế (formatted: 270.000)
  {BANG_CHU}            - Bằng chữ (Hai trăm bảy mươi nghìn đồng.)
  {NGAY_HOP_DONG}       - ngày DD tháng MM năm YYYY

CCCD images: Insert at paragraphs containing "Mặt trước CCCD" / "Mặt sau CCCD".
"""

import os
import re
import copy
import tempfile
from datetime import datetime
from typing import Dict, Optional, List
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)

# Template path - local fallback
TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")
TEMPLATE_HDKOC = os.path.join(TEMPLATE_DIR, "Mau_hop_dong_KOC.docx")

# Template name mapping: Lark "Template" field value → Drive search name
TEMPLATE_MAP = {
    "HDKOC": "HDKOC",
}


# ╔════════════════════════════════════════════════════════════════╗
# ║              FORMAT CHUẨN HỢP ĐỒNG VIỆT NAM                   ║
# ╚════════════════════════════════════════════════════════════════╝

MARGIN_TOP = Cm(2)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(1.5)

LINE_SPACING = 1.15
SPACE_AFTER_NORMAL = Pt(3)
SPACE_AFTER_HEADING = Pt(6)
SPACE_AFTER_MAX = Pt(8)
SPACE_BEFORE_HEADING = Pt(6)
BIG_SPACE_THRESHOLD = Pt(10)
HEADING_KEYWORDS = ["ĐIỀU", "CỘNG HOÀ", "HỢP ĐỒNG", "ĐẠI DIỆN"]


# ╔════════════════════════════════════════════════════════════════╗
# ║              PLACEHOLDER FIND & REPLACE ENGINE                  ║
# ╚════════════════════════════════════════════════════════════════╝

def _replace_in_paragraph(paragraph, placeholder: str, replacement: str) -> bool:
    """
    Replace {PLACEHOLDER} in a paragraph, handling cross-run cases.
    
    Word often splits text across multiple runs (e.g. different formatting),
    so {THUONG_HIEU} might be split as "{THUONG" in run[0] and "_HIEU}" in run[1].
    This function handles all cases.
    
    Returns True if replacement was made.
    """
    runs = paragraph.runs
    if not runs:
        return False
    
    full_text = ''.join(run.text for run in runs)
    if placeholder not in full_text:
        return False
    
    # Simple case: placeholder is entirely within a single run
    for run in runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            return True
    
    # Complex case: placeholder spans multiple runs
    start_pos = full_text.index(placeholder)
    end_pos = start_pos + len(placeholder)
    
    char_count = 0
    found_start = False
    
    for run in runs:
        run_start = char_count
        run_end = char_count + len(run.text)
        
        if not found_start and run_start <= start_pos < run_end:
            found_start = True
            before = run.text[:start_pos - run_start]
            
            if run_end >= end_pos:
                after = run.text[end_pos - run_start:]
                run.text = before + replacement + after
                return True
            else:
                run.text = before + replacement
                
        elif found_start and run_end <= end_pos:
            run.text = ''
            
        elif found_start and run_start < end_pos <= run_end:
            after = run.text[end_pos - run_start:]
            run.text = after
            return True
        
        char_count = run_end
    
    return True


def _replace_all(doc: Document, replacements: Dict[str, str]) -> Dict[str, int]:
    """
    Replace all placeholders throughout the entire document.
    Searches: paragraphs, table cells, headers, footers.
    
    Returns dict of placeholder → count of replacements made.
    """
    counts = {k: 0 for k in replacements}
    
    for placeholder, replacement in replacements.items():
        if not replacement and replacement != "":
            replacement = ""
        replacement = str(replacement)
        
        # Search in paragraphs
        for para in doc.paragraphs:
            while placeholder in ''.join(r.text for r in para.runs):
                if _replace_in_paragraph(para, placeholder, replacement):
                    counts[placeholder] += 1
                else:
                    break
        
        # Search in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        while placeholder in ''.join(r.text for r in para.runs):
                            if _replace_in_paragraph(para, placeholder, replacement):
                                counts[placeholder] += 1
                            else:
                                break
        
        # Search in headers/footers
        for section in doc.sections:
            for hf in [section.header, section.footer]:
                if hf and hf.paragraphs:
                    for para in hf.paragraphs:
                        while placeholder in ''.join(r.text for r in para.runs):
                            if _replace_in_paragraph(para, placeholder, replacement):
                                counts[placeholder] += 1
                            else:
                                break
    
    return counts


# ╔════════════════════════════════════════════════════════════════╗
# ║                    FORMAT HELPERS                               ║
# ╚════════════════════════════════════════════════════════════════╝

def _is_heading_paragraph(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    for kw in HEADING_KEYWORDS:
        if stripped.startswith(kw):
            return True
    return False


def _normalize_formatting(doc: Document):
    """Chuẩn hóa formatting toàn bộ document theo tiêu chuẩn hợp đồng VN."""
    
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
    
    for i, para in enumerate(doc.paragraphs):
        pf = para.paragraph_format
        text = para.text.strip()
        
        try:
            if pf.first_line_indent is not None and pf.first_line_indent < 0:
                pf.first_line_indent = 0
        except (ValueError, TypeError):
            pass
        
        try:
            if pf.space_after is not None and pf.space_after > BIG_SPACE_THRESHOLD:
                if _is_heading_paragraph(text):
                    pf.space_after = SPACE_AFTER_HEADING
                else:
                    pf.space_after = SPACE_AFTER_NORMAL
        except (ValueError, TypeError):
            pass
        
        try:
            if pf.line_spacing is not None and isinstance(pf.line_spacing, float):
                if 1.0 <= pf.line_spacing <= 1.2:
                    pf.line_spacing = LINE_SPACING
        except (ValueError, TypeError):
            pass
        
        if "Hà Nội, ngày" in text:
            clean_date = text.replace("\t", "").strip()
            if para.runs:
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = clean_date
            pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf.first_line_indent = 0
            pf.space_after = SPACE_AFTER_HEADING
        
        if text.startswith("(Về việc"):
            pf.space_after = SPACE_AFTER_HEADING
        
        if text.startswith("ĐIỀU"):
            pf.space_before = SPACE_BEFORE_HEADING
            pf.space_after = SPACE_AFTER_HEADING
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pf = para.paragraph_format
                    try:
                        if pf.first_line_indent is not None and pf.first_line_indent < 0:
                            pf.first_line_indent = 0
                    except (ValueError, TypeError):
                        pass
                    try:
                        if pf.space_after is not None and pf.space_after > BIG_SPACE_THRESHOLD:
                            pf.space_after = Pt(2)
                    except (ValueError, TypeError):
                        pass
                    try:
                        if pf.line_spacing is not None and isinstance(pf.line_spacing, float):
                            if 1.0 <= pf.line_spacing <= 1.2:
                                pf.line_spacing = LINE_SPACING
                    except (ValueError, TypeError):
                        pass


def _format_currency_vn(amount) -> str:
    """Format number as Vietnamese currency: 11700000 → '11.700.000'"""
    if not amount and amount != 0:
        return ""
    try:
        s = str(amount).replace(",", "").strip()
        # Single dot = decimal point (e.g. Lark returns float 11111111.0)
        # Multiple dots = VN thousands separator (e.g. "11.111.111")
        if s.count(".") == 1:
            num = int(float(s))
        else:
            num = int(float(s.replace(".", ""))) if s else 0
        if num == 0:
            return "0"
        return f"{num:,}".replace(",", ".")
    except (ValueError, TypeError):
        return str(amount)


def _format_date_field(value) -> str:
    """Convert Lark date field to dd/mm/yyyy string."""
    if not value:
        return ""
    s = str(value).strip()
    if "/" in s and len(s) <= 10:
        return s
    try:
        ts = float(s)
        if ts > 1e12:
            ts = ts / 1000
        dt = datetime.fromtimestamp(ts)
        return dt.strftime("%d/%m/%Y")
    except (ValueError, TypeError, OSError):
        pass
    return s


def _format_percent(value) -> str:
    """Convert percentage: 0.05 → '5', '5%' → '5', 0.1 → '10'"""
    if not value and value != 0:
        return ""
    s = str(value).replace("%", "").strip()
    try:
        num = float(s)
        if num < 1:
            num = num * 100
        return f"{num:g}"
    except (ValueError, TypeError):
        return s


def _format_date_vn(date_value) -> str:
    """Convert date value to Vietnamese format dd/MM/yyyy."""
    if not date_value:
        return ""
    if isinstance(date_value, (int, float)):
        dt = datetime.fromtimestamp(date_value / 1000)
        return dt.strftime("%d/%m/%Y")
    if isinstance(date_value, datetime):
        return date_value.strftime("%d/%m/%Y")
    if isinstance(date_value, str):
        if len(date_value) == 10 and date_value[2] == "/" and date_value[5] == "/":
            return date_value
        try:
            dt = datetime.fromisoformat(date_value.replace("Z", "+00:00"))
            return dt.strftime("%d/%m/%Y")
        except ValueError:
            pass
        try:
            dt = datetime.strptime(date_value, "%Y-%m-%d")
            return dt.strftime("%d/%m/%Y")
        except ValueError:
            pass
    return str(date_value)


def _get_current_date_vn() -> str:
    now = datetime.now()
    return f"ngày {now.day:02d} tháng {now.month:02d} năm {now.year}"


def _number_to_vietnamese_words(number) -> str:
    """Chuyển số thành chữ tiếng Việt."""
    try:
        s = str(number).replace(",", "").strip()
        if s.count(".") == 1:
            n = int(float(s))
        else:
            n = int(float(s.replace(".", ""))) if s else 0
    except (ValueError, TypeError):
        return ""
    
    if n == 0:
        return "Không đồng"
    
    ones = ["", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
    
    def _read_group_3(num):
        if num == 0:
            return ""
        h = num // 100
        t = (num % 100) // 10
        u = num % 10
        parts = []
        if h > 0:
            parts.append(f"{ones[h]} trăm")
            if t == 0 and u > 0:
                parts.append("lẻ")
        if t > 1:
            parts.append(f"{ones[t]} mươi")
            if u == 1:
                parts.append("mốt")
            elif u == 4:
                parts.append("tư")
            elif u == 5:
                parts.append("lăm")
            elif u > 0:
                parts.append(ones[u])
        elif t == 1:
            parts.append("mười")
            if u == 5:
                parts.append("lăm")
            elif u > 0:
                parts.append(ones[u])
        elif t == 0 and h > 0 and u > 0:
            parts.append(ones[u])
        elif t == 0 and h == 0 and u > 0:
            parts.append(ones[u])
        return " ".join(parts)
    
    units = ["", "nghìn", "triệu", "tỷ"]
    groups = []
    temp = n
    while temp > 0:
        groups.append(temp % 1000)
        temp //= 1000
    
    result_parts = []
    for i in range(len(groups) - 1, -1, -1):
        if groups[i] > 0:
            text = _read_group_3(groups[i])
            if text:
                if i < len(units):
                    result_parts.append(f"{text} {units[i]}".strip())
                else:
                    result_parts.append(text)
    
    result = " ".join(result_parts).strip()
    result = result[0].upper() + result[1:] if result else ""
    return f"{result} đồng."


# ╔════════════════════════════════════════════════════════════════╗
# ║                    CCCD IMAGE INSERTION                         ║
# ╚════════════════════════════════════════════════════════════════╝

def _insert_cccd_image(doc: Document, keyword: str, image_path: str):
    """Insert CCCD image after the paragraph containing the keyword."""
    try:
        from docx.shared import Cm
        from PIL import Image as PILImage
        import io, traceback
        
        target_idx = None
        for i, para in enumerate(doc.paragraphs):
            if keyword in para.text:
                target_idx = i
                break
        
        if target_idx is None:
            print(f"⚠️ CCCD keyword '{keyword}' not found in document")
            return
        
        img = PILImage.open(image_path)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=90)
        buf.seek(0)
        
        print(f"📐 CCCD image {image_path}: {img.size[0]}x{img.size[1]}, converted to JPEG ({buf.getbuffer().nbytes} bytes)")
        
        target_para = None
        for idx in range(target_idx + 1, min(target_idx + 4, len(doc.paragraphs))):
            p = doc.paragraphs[idx]
            if not p.text.strip():
                target_para = p
                break
        
        if target_para is None:
            target_para = doc.paragraphs[target_idx]
        
        for run in target_para.runs:
            run.text = ""
        
        run = target_para.add_run()
        run.add_picture(buf, width=Cm(14))
        print(f"✅ Inserted CCCD image at '{keyword}'")
    except Exception as e:
        print(f"⚠️ Failed to insert CCCD image '{keyword}': {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()


# ╔════════════════════════════════════════════════════════════════╗
# ║                    MAIN GENERATE FUNCTION                       ║
# ╚════════════════════════════════════════════════════════════════╝

def generate_contract(data: Dict, template_path: str = None, output_path: str = None) -> str:
    """
    Generate a KOC contract by replacing placeholders in Word template.
    Works with ANY template structure - no hardcoded indices.
    """
    template = template_path
    if not template:
        template_key = data.get("template", "HDKOC")
        drive_name = TEMPLATE_MAP.get(template_key, template_key)
        try:
            from google_drive_client import get_template_path
            template = get_template_path(drive_name, fallback_path=TEMPLATE_HDKOC)
        except Exception as e:
            print(f"⚠️ Drive template fetch failed, using local: {e}")
            template = TEMPLATE_HDKOC
    
    if not os.path.exists(template):
        raise FileNotFoundError(f"Template not found: {template}")
    
    doc = Document(template)
    
    # Extract fields
    ho_ten = data.get("ho_ten", "")
    id_koc = data.get("id_koc", "")
    dia_chi = data.get("dia_chi", "")
    mst = data.get("mst", "")
    sdt = data.get("sdt", "")
    gmail = data.get("gmail", "")
    cccd = data.get("cccd", "")
    cccd_ngay_cap = _format_date_vn(data.get("cccd_ngay_cap", ""))
    cccd_noi_cap = data.get("cccd_noi_cap", "")
    stk = data.get("stk", "")
    
    thuong_hieu = data.get("thuong_hieu", "")
    ngay_du_kien_air = _format_date_field(data.get("ngay_du_kien_air", ""))
    hoa_hong_tu_nhien = _format_percent(data.get("hoa_hong_tu_nhien", ""))
    hoa_hong_chay_ads = _format_percent(data.get("hoa_hong_chay_ads", ""))
    
    thanh_tien = data.get("thanh_tien", "")
    chi_phi = data.get("chi_phi", "")
    so_luong_clip = data.get("so_luong_clip", "")
    thue_tncn = data.get("thue_tncn", "")
    tong_gia_tri_sau_thue = data.get("tong_gia_tri_sau_thue", "")
    
    # Calculate chi_phi if not provided
    if not chi_phi and thanh_tien and so_luong_clip:
        try:
            chi_phi = str(int(float(str(thanh_tien)) / float(str(so_luong_clip))))
            print(f"💰 Chi phí calculated: {thanh_tien} / {so_luong_clip} = {chi_phi}")
        except (ValueError, ZeroDivisionError):
            pass
    
    # Format so_luong_clip as integer
    if so_luong_clip:
        try:
            so_luong_clip = str(int(float(str(so_luong_clip))))
        except (ValueError, TypeError):
            pass
    
    # Calculate bằng chữ
    bang_chu = ""
    if tong_gia_tri_sau_thue:
        bang_chu = _number_to_vietnamese_words(tong_gia_tri_sau_thue)
    
    print(f"📋 Template structure: {len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs")
    
    # ═══════════════════════════════════════════════════════
    # BUILD REPLACEMENT MAP
    # ═══════════════════════════════════════════════════════
    replacements = {
        "{HO_TEN}": str(ho_ten),
        "{DIA_CHI}": str(dia_chi),
        "{MST}": str(mst),
        "{SDT}": str(sdt),
        "{GMAIL}": str(gmail),
        "{CCCD}": str(cccd),
        "{CCCD_NGAY_CAP}": str(cccd_ngay_cap),
        "{CCCD_NOI_CAP}": str(cccd_noi_cap),
        "{STK}": str(stk),
        "{ID_KOC}": str(id_koc),
        "{THUONG_HIEU}": str(thuong_hieu),
        "{NGAY_DU_KIEN_AIR}": str(ngay_du_kien_air),
        "{HOA_HONG_TU_NHIEN}": str(hoa_hong_tu_nhien),
        "{HOA_HONG_CHAY_ADS}": str(hoa_hong_chay_ads),
        "{SO_LUONG_CLIP}": str(so_luong_clip),
        "{CHI_PHI}": _format_currency_vn(chi_phi),
        "{THANH_TIEN}": _format_currency_vn(thanh_tien),
        "{THUE_TNCN}": _format_currency_vn(thue_tncn),
        "{TONG_SAU_THUE}": _format_currency_vn(tong_gia_tri_sau_thue),
        "{BANG_CHU}": str(bang_chu),
        "{NGAY_HOP_DONG}": _get_current_date_vn(),
    }
    
    # ═══════════════════════════════════════════════════════
    # REPLACE ALL PLACEHOLDERS
    # ═══════════════════════════════════════════════════════
    counts = _replace_all(doc, replacements)
    
    found = {k: v for k, v in counts.items() if v > 0}
    not_found = [k for k, v in counts.items() if v == 0 and replacements[k]]
    print(f"✅ Replaced {sum(counts.values())} placeholders: {found}")
    if not_found:
        print(f"⚠️ Placeholders not found in template: {not_found}")
    
    # ═══════════════════════════════════════════════════════
    # INSERT CCCD IMAGES (by keyword search, not index)
    # ═══════════════════════════════════════════════════════
    cccd_truoc_path = data.get("cccd_truoc_path", "")
    cccd_sau_path = data.get("cccd_sau_path", "")
    
    if cccd_truoc_path and os.path.exists(cccd_truoc_path):
        fsize = os.path.getsize(cccd_truoc_path)
        if fsize > 100:
            _insert_cccd_image(doc, "Mặt trước CCCD", cccd_truoc_path)
        else:
            print(f"⚠️ CCCD front file too small: {fsize} bytes")
    
    if cccd_sau_path and os.path.exists(cccd_sau_path):
        fsize = os.path.getsize(cccd_sau_path)
        if fsize > 100:
            _insert_cccd_image(doc, "Mặt sau CCCD", cccd_sau_path)
        else:
            print(f"⚠️ CCCD back file too small: {fsize} bytes")
    
    # ═══════════════════════════════════════════════════════
    # NORMALIZE FORMATTING
    # ═══════════════════════════════════════════════════════
    _normalize_formatting(doc)
    
    # ═══════════════════════════════════════════════════════
    # SAVE OUTPUT
    # ═══════════════════════════════════════════════════════
    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path
    
    safe_name = "".join(c for c in ho_ten if c.isalnum() or c in " _-").strip().replace(" ", "_")
    filename = f"{id_koc}_{datetime.now().strftime('%d-%m-%Y')}.docx"
    
    tmp_dir = tempfile.mkdtemp(prefix="contract_")
    output = os.path.join(tmp_dir, filename)
    doc.save(output)
    
    logger.info(f"✅ Contract generated: {output}")
    return output


def parse_lark_record_to_contract_data(fields: Dict) -> Dict:
    """Convert Lark Base record fields to contract_generator data format."""
    sdt = fields.get("SDT Bên B", "")
    if isinstance(sdt, dict):
        sdt = sdt.get("value", sdt.get("text", ""))
    
    return {
        "id_koc": fields.get("ID KOC", ""),
        "ho_ten": fields.get("Họ và Tên Bên B", ""),
        "dia_chi": fields.get("Địa chỉ Bên B", ""),
        "mst": fields.get("MST Bên B", ""),
        "sdt": str(sdt),
        "cccd": fields.get("CCCD Bên B", ""),
        "cccd_ngay_cap": fields.get("CCCD Ngày Cấp", ""),
        "cccd_noi_cap": fields.get("CCCD Nơi Cấp", ""),
        "gmail": fields.get("Gmail Bên B", ""),
        "stk": fields.get("STK bên B", ""),
        # Payment
        "chi_phi": fields.get("Chi phí", ""),
        "thanh_tien": fields.get("Thành Tiền", fields.get("Thành Tiên", "")),
        "so_luong_clip": fields.get("Số lượng clip", ""),
        "thue_tncn": fields.get("Thuế TNCN", ""),
        "tong_gia_tri_sau_thue": fields.get("Tổng giá trị sau thuế", ""),
        # Paragraph fields
        "thuong_hieu": fields.get("Thương hiệu", ""),
        "ngay_du_kien_air": fields.get("Ngày dự kiến air", ""),
        "hoa_hong_tu_nhien": fields.get("% Hoa hồng tự nhiên", fields.get("Hoa hồng tự nhiên", "")),
        "hoa_hong_chay_ads": fields.get("% Hoa hồng chạy ads", fields.get("Hoa hồng chạy ads", "")),
        # CCCD images
        "cccd_truoc_path": fields.get("cccd_truoc_path", ""),
        "cccd_sau_path": fields.get("cccd_sau_path", ""),
        # Template
        "template": fields.get("Template", "HDKOC"),
    }
